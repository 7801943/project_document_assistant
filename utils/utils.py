import psutil # 用于IPv6地址处理
import socket # 用于 socket.AF_INET6
import time # 用于时间处理
import hashlib # 用于MD5哈希计算
from pathlib import Path # 用于路径操作
from typing import Optional, Tuple # 用于类型提示
import httpx # 导入 httpx
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from loguru import logger # 从 loguru 导入 logger
from config import settings # 导入配置 settings

async def check_embedding_service_health(client: httpx.AsyncClient) -> bool:
    """
    检查嵌入模型服务的健康状况。
    通过请求模型的形式来验证服务端点是否可用
    """
    # 确保基础 URL 没有尾随斜杠，然后附加 '/models'
    api_url = f"{str(settings.EMBEDDING_API_URL).rstrip('/')}/models"
    headers = {
        "Authorization": f"Bearer {settings.EMBEDDING_APIKEY.get_secret_value()}"
    }
    try:
        logger.info(f"正在检查嵌入模型服务，目标URL: {api_url}")
        # 为健康检查设置一个较短的超时时间，例如5秒
        response = await client.get(api_url, headers=headers, timeout=5.0)
        
        if response.status_code == 200:
            # logger.info("嵌入模型服务健康检查成功，服务端点可用。")
            # 可选: 更详细地检查响应内容
            models = response.json()
            model_names = [model.get('id') for model in models.get('data', [])]
            logger.info(f"嵌入模型服务可用，模型列表: {model_names}")
            return True
        else:
            logger.warning(f"嵌入模型服务健康检查返回非200状态码: {response.status_code}，响应: {response.text}")
            return False
    except httpx.RequestError as e:
        # 捕获所有 httpx 请求相关的错误 (例如: 连接错误, 超时)
        logger.error(f"无法连接到嵌入模型服务或请求超时: {e}")
        return False
    except Exception as e:
        # 捕获其他任何意外错误
        logger.error(f"检查嵌入模型服务时发生未知错误: {e}", exc_info=True)
        return False

def get_host_ipv6_addr(interface: Optional[str] = None) -> str:
    """
    获取指定网络接口的非链接本地IPv6地址。
    如果未提供接口，则使用配置中的 settings.SERVER_INTERFACE。
    """
    # 如果函数调用时未指定 interface，则使用配置文件中的值
    if interface is None:
        interface_to_use = settings.SERVER_INTERFACE
    else:
        interface_to_use = interface

    try:
        if_addrs = psutil.net_if_addrs()
        if interface_to_use not in if_addrs:
            logger.info(f"未获取到接口 {interface_to_use} 的 IP 信息")
            return ""
        for addr in if_addrs[interface_to_use]:
            if addr.family == socket.AF_INET6 and not addr.address.startswith('fe80::'):
                ipv6_addr = addr.address.split('%')[0] # 移除可能存在的 scope id (如 %eth0)
                return ipv6_addr
        logger.info(f"接口 {interface_to_use} 上未找到合适的 IPv6 地址。")
        return ""
    except Exception as e:
        logger.error(f"获取接口 {interface_to_use} 的 IPv6 地址失败: {e}")
        return ""

def get_current_time()-> Tuple[float, str]:
    """
    获取当前时间戳和格式化后的时间字符串。
    返回:
        一个元组，包含 (时间戳_float, "YYYY-MM-DD HH:MM:SS"格式的字符串)
    """
    current_timestamp = time.time()
    local_time_struct = time.localtime(current_timestamp)
    formatted_string = time.strftime("%Y-%m-%d %H:%M:%S", local_time_struct)
    return (current_timestamp, formatted_string)

def calculate_md5(file_path: Path) -> Optional[str]:
    """
    计算文件的MD5哈希值。
    参数:
        file_path: 文件的 Path 对象。
    返回:
        文件的MD5哈希字符串，如果文件不存在或读取错误则返回 None。
    """
    if not file_path.is_file():
        logger.warning(f"请求计算MD5的文件不是一个有效文件: {file_path}")
        return None
    hash_md5 = hashlib.md5()
    try:
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""): # 每次读取 4KB
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    except IOError as e:
        logger.error(f"无法读取文件 {file_path} 以计算MD5: {e}")
        return None
    except Exception as e: # 捕获其他潜在错误
        logger.error(f"计算文件 {file_path} MD5时发生未知错误: {e}")
        return None


# def remove_whitespace_from_docx(path: str):
#     '''
#     去除docx中的空格
#     '''
#     doc = Document(path)
#     for para in doc.paragraphs:
#         for run in para.runs:
#             run.text = run.text.replace('\u00A0', '')  # 非断行空格
#             run.text = run.text.replace(' ', '')       # 普通空格
#             run.text = run.text.replace('\t', '')      # 制表符
#             run.text = run.text.replace('\n', '')      # 显式换行符

#     doc.save(path)

# def remove_empty_paragraphs_from_docx(path: str):
#     doc = Document(path)

#     # 新建文档
#     new_doc = Document()

#     for para in doc.paragraphs:
#         if para.text.strip():  # 只复制非空段落
#             new_para = new_doc.add_paragraph()
#             for run in para.runs:
#                 new_run = new_para.add_run(run.text)
#                 # 保留基本格式
#                 new_run.bold = run.bold
#                 new_run.italic = run.italic
#                 new_run.underline = run.underline
#                 new_run.font.name = run.font.name

#     new_doc.save(path)

def remove_empty_paragraphs(doc_path):
    """
    从 .docx 文件中移除空段落和仅包含空白字符的段落。
    此函数会同时处理文档正文和表格单元格中的段落。

    :param doc_path: 原始 .docx 文件的路径。
    :param new_doc_path: (可选) 清理后新文档的保存路径。
                         如果为 None，则会自动生成一个新文件名，
                         例如 "原文件名_cleaned.docx"。
    """
    
    # 内部辅助函数，用于删除段落的 XML 元素
    def _delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        # 将段落对象的内部引用设为 None，是一种好的编程习惯
        paragraph._p = paragraph._element = None

    # 检查文件是否存在
    # if not os.path.exists(doc_path):
    #     print(f"错误：文件 '{doc_path}' 不存在。")
    #     return

    # 如果未提供新路径，则自动生成一个
    # if new_doc_path is None:
    #     base, ext = os.path.splitext(doc_path)
    #     new_doc_path = f"{base}_cleaned{ext}"

    doc = Document(doc_path)

    # --- 步骤 1: 识别所有待删除的段落 ---
    # 创建一个列表来收集所有需要删除的段落对象
    paragraphs_to_delete = []
    
    # 遍历正文段落
    for paragraph in doc.paragraphs:
        # 如果段落去除首尾空白后长度为0，则判定为空
        if len(paragraph.text.strip()) == 0:
            paragraphs_to_delete.append(paragraph)

    # --- 步骤 2: 一次性删除所有已识别的段落 ---
    # 在一个单独的循环中执行删除操作，避免在遍历时修改集合
    for p in paragraphs_to_delete:
        _delete_paragraph(p)

    # --- 步骤 3: 保存清理后的文档 ---
    try:
        doc.save(doc_path)
        # print(f"成功！已移除 {len(paragraphs_to_delete)} 个空段落。")
        # print(f"清理后的文档已保存至：{new_doc_path}")
    except Exception as e:
        print(f"保存文件时出错：{e}")